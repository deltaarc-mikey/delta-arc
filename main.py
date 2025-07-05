import streamlit as st
import pandas as pd
import os
from openai import OpenAI
from docx import Document

# Load OpenAI API key from environment
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

st.set_page_config(page_title="Strategy Replay Mode + GPT Summary", layout="wide")
st.title("📘 Strategy Replay Mode + Batch Backtest Loop")

# File Upload
st.subheader("📅 Upload Historical Price Data (CSV)")
price_file = st.file_uploader("Upload your historical CSV", type="csv")
if price_file:
    df = pd.read_csv(price_file)
    if 'Date' in df.columns:
        df['Date'] = pd.to_datetime(df['Date'])
        st.success("CSV uploaded successfully.")
        st.dataframe(df.head())
    else:
        st.error("Missing 'Date' column in CSV.")

st.subheader("📑 Upload Trade History (CSV)")
trade_file = st.file_uploader("Upload your trades CSV", type="csv", key="trades")
if trade_file:
    trades = pd.read_csv(trade_file)
    st.success("Trade CSV uploaded.")
    st.dataframe(trades.head())

# Run Summary Comparison
st.subheader("🧠 Generate GPT-4 Summary")
input_text = st.text_area("Paste input for GPT-4 summary (e.g. trade batch details)", height=180)

if st.button("Run GPT-4 Summary"):
    try:
        with st.spinner("Summarizing with GPT-4..."):
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a professional financial analyst. Provide a clear, expert-level evaluation of the following trade data."},
                    {"role": "user", "content": input_text}
                ]
            )
            gpt_summary = response.choices[0].message.content
            st.markdown("### 🤖 GPT-4 Summary")
            st.write(gpt_summary)

            # Save to DOCX
            doc = Document()
            doc.add_heading("GPT-4 Summary", 0)
            doc.add_paragraph(gpt_summary)
            doc_path = "gpt_summary.docx"
            doc.save(doc_path)
            with open(doc_path, "rb") as file:
                st.download_button("📥 Download GPT Summary (.docx)", file, file_name=doc_path)

    except Exception as e:
        st.error(f"Error generating GPT summary: {str(e)}")

# Claude summary input (manual)
st.subheader("📝 Claude Summary (Paste)")
claude_summary = st.text_area("Paste Claude Summary", height=150)

# Side-by-side comparison
if input_text and claude_summary:
    st.subheader("📊 Claude vs GPT Summary Comparison")
    with st.expander("📋 Summary Comparison"):
        st.markdown("#### Claude Summary:")
        st.markdown(claude_summary)

        if 'gpt_summary' in locals():
            st.markdown("#### GPT Summary:")
            st.markdown(gpt_summary)

            st.markdown("#### 📌 Verdict:")
            st.success("✅ Both GPT and Claude provided insight. Use this comparison for strategy refinement or case study archiving.")

# Instruction
st.info("👉 Upload both historical price CSV and trade CSV to begin full backtest. Use GPT and Claude summaries to archive insights or identify future entry conditions.")
